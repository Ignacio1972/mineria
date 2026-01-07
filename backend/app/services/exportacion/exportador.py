"""
Exportador de Informes de Prefactibilidad

Genera informes en formato PDF y DOCX a partir de los datos
del análisis de prefactibilidad ambiental.
"""

import logging
import io
from enum import Enum
from typing import Any, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class FormatoExportacion(str, Enum):
    """Formatos de exportación disponibles"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"


class ExportadorInformes:
    """
    Exportador de informes de prefactibilidad ambiental.

    Soporta exportación a:
    - PDF: Documento formal con formato profesional
    - DOCX: Documento editable de Microsoft Word
    - TXT: Texto plano
    - HTML: Documento HTML para visualización web
    """

    # Colores corporativos
    COLOR_PRIMARIO = (0.2, 0.4, 0.6)  # Azul
    COLOR_CRITICO = (0.8, 0.2, 0.2)   # Rojo
    COLOR_ALTO = (0.9, 0.5, 0.1)      # Naranja
    COLOR_MEDIO = (0.9, 0.8, 0.1)     # Amarillo
    COLOR_BAJO = (0.2, 0.7, 0.3)      # Verde

    def __init__(self):
        logger.info("ExportadorInformes inicializado")

    def exportar(
        self,
        informe: dict[str, Any],
        formato: FormatoExportacion = FormatoExportacion.PDF,
    ) -> bytes:
        """
        Exporta un informe al formato especificado.

        Args:
            informe: Diccionario con los datos del informe
            formato: Formato de exportación (debe ser FormatoExportacion enum)

        Returns:
            Bytes del documento generado

        Raises:
            ValueError: Si el formato no es válido
        """
        # Validar que formato sea del tipo correcto
        if not isinstance(formato, FormatoExportacion):
            raise ValueError(f"Formato no válido: {formato}. Debe usar FormatoExportacion enum.")

        logger.info(f"Exportando informe a formato: {formato.value}")

        if formato == FormatoExportacion.PDF:
            return self._exportar_pdf(informe)
        elif formato == FormatoExportacion.DOCX:
            return self._exportar_docx(informe)
        elif formato == FormatoExportacion.TXT:
            return self._exportar_txt(informe)
        elif formato == FormatoExportacion.HTML:
            return self._exportar_html(informe)
        else:
            raise ValueError(f"Formato no soportado: {formato}")

    def _exportar_pdf(self, informe: dict[str, Any]) -> bytes:
        """Genera el informe en formato PDF."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, ListFlowable, ListItem
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=LETTER,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
        )

        # Estilos
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='TituloInforme',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1a4d7c'),
        ))
        styles.add(ParagraphStyle(
            name='Subtitulo',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.HexColor('#2d6da3'),
        ))
        styles.add(ParagraphStyle(
            name='SeccionTitulo',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor('#3d7db3'),
        ))
        styles.add(ParagraphStyle(
            name='Cuerpo',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            name='Alerta',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=6,
        ))

        # Contenido
        story = []

        # Título
        datos_proyecto = informe.get('datos_proyecto', {})
        story.append(Paragraph(
            "INFORME DE PREFACTIBILIDAD AMBIENTAL",
            styles['TituloInforme']
        ))
        story.append(Paragraph(
            datos_proyecto.get('nombre', 'Proyecto Sin Nombre'),
            styles['Subtitulo']
        ))
        story.append(Spacer(1, 20))

        # Información del informe
        fecha = informe.get('fecha_generacion', datetime.now().isoformat())
        if isinstance(fecha, str):
            try:
                fecha = datetime.fromisoformat(fecha.replace('Z', '+00:00'))
                fecha_str = fecha.strftime('%d/%m/%Y %H:%M')
            except ValueError:
                fecha_str = fecha
        else:
            fecha_str = fecha.strftime('%d/%m/%Y %H:%M')

        info_data = [
            ['ID Informe:', informe.get('id', 'N/D')],
            ['Fecha:', fecha_str],
            ['Vía Recomendada:', informe.get('clasificacion_seia', {}).get('via_ingreso_recomendada', 'N/D')],
            ['Confianza:', f"{informe.get('clasificacion_seia', {}).get('confianza', 0) * 100:.0f}%"],
        ]
        info_table = Table(info_data, colWidths=[1.5*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1a4d7c')),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))

        # Secciones del informe
        secciones = informe.get('secciones', [])
        if isinstance(secciones, list):
            for seccion in secciones:
                if isinstance(seccion, dict):
                    story.append(Paragraph(
                        seccion.get('titulo', 'Sin Título'),
                        styles['SeccionTitulo']
                    ))
                    contenido = seccion.get('contenido', '')
                    # Dividir en párrafos
                    for parrafo in contenido.split('\n\n'):
                        if parrafo.strip():
                            # Manejar listas con viñetas
                            if parrafo.strip().startswith('- '):
                                items = [p.strip()[2:] for p in parrafo.split('\n') if p.strip().startswith('- ')]
                                for item in items:
                                    story.append(Paragraph(f"• {item}", styles['Alerta']))
                            else:
                                story.append(Paragraph(parrafo.strip(), styles['Cuerpo']))
                    story.append(Spacer(1, 10))

        # Alertas
        alertas = informe.get('alertas', [])
        if alertas:
            story.append(PageBreak())
            story.append(Paragraph("ALERTAS AMBIENTALES", styles['Subtitulo']))

            for alerta in alertas[:10]:  # Limitar a 10 alertas
                nivel = alerta.get('nivel', 'INFO')
                color = {
                    'CRITICA': colors.HexColor('#cc3333'),
                    'ALTA': colors.HexColor('#e67300'),
                    'MEDIA': colors.HexColor('#cccc00'),
                    'BAJA': colors.HexColor('#33cc33'),
                }.get(nivel, colors.black)

                story.append(Paragraph(
                    f"<font color='{color.hexval()}'>[{nivel}]</font> {alerta.get('titulo', 'N/D')}",
                    styles['SeccionTitulo']
                ))
                story.append(Paragraph(alerta.get('descripcion', ''), styles['Cuerpo']))

        # Normativa citada
        normativa = informe.get('normativa_citada', [])
        if normativa:
            story.append(Spacer(1, 20))
            story.append(Paragraph("NORMATIVA CITADA", styles['Subtitulo']))
            for norm in normativa[:15]:
                ref = norm.get('referencia', '') or norm.get('documento', '')
                if ref:
                    story.append(Paragraph(f"• {ref}", styles['Alerta']))

        # Disclaimer
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            "<i>Este informe es un análisis preliminar de prefactibilidad ambiental. "
            "No reemplaza la evaluación formal del SEA ni la opinión de expertos especializados.</i>",
            styles['Cuerpo']
        ))

        # Generar PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _exportar_docx(self, informe: dict[str, Any]) -> bytes:
        """Genera el informe en formato DOCX."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Estilos personalizados
        styles = doc.styles

        # Título
        titulo = doc.add_heading('INFORME DE PREFACTIBILIDAD AMBIENTAL', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Nombre del proyecto
        datos_proyecto = informe.get('datos_proyecto', {})
        subtitulo = doc.add_heading(datos_proyecto.get('nombre', 'Proyecto Sin Nombre'), 1)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Información del informe
        fecha = informe.get('fecha_generacion', datetime.now().isoformat())
        if isinstance(fecha, str):
            try:
                fecha = datetime.fromisoformat(fecha.replace('Z', '+00:00'))
                fecha_str = fecha.strftime('%d/%m/%Y %H:%M')
            except ValueError:
                fecha_str = fecha
        else:
            fecha_str = fecha.strftime('%d/%m/%Y %H:%M')

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        info_data = [
            ('ID Informe:', informe.get('id', 'N/D')),
            ('Fecha:', fecha_str),
            ('Vía Recomendada:', informe.get('clasificacion_seia', {}).get('via_ingreso_recomendada', 'N/D')),
            ('Confianza:', f"{informe.get('clasificacion_seia', {}).get('confianza', 0) * 100:.0f}%"),
        ]
        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Secciones del informe
        secciones = informe.get('secciones', [])
        for seccion in secciones:
            if isinstance(seccion, dict):
                doc.add_heading(seccion.get('titulo', 'Sin Título'), 2)
                contenido = seccion.get('contenido', '')
                for parrafo in contenido.split('\n\n'):
                    if parrafo.strip():
                        if parrafo.strip().startswith('- '):
                            for line in parrafo.split('\n'):
                                if line.strip().startswith('- '):
                                    p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif parrafo.strip().startswith('**'):
                            # Texto en negrita
                            p = doc.add_paragraph()
                            text = parrafo.strip().replace('**', '')
                            run = p.add_run(text)
                            run.bold = True
                        else:
                            doc.add_paragraph(parrafo.strip())

        # Alertas
        alertas = informe.get('alertas', [])
        if alertas:
            doc.add_page_break()
            doc.add_heading('ALERTAS AMBIENTALES', 1)

            for alerta in alertas[:10]:
                nivel = alerta.get('nivel', 'INFO')
                titulo_alerta = f"[{nivel}] {alerta.get('titulo', 'N/D')}"
                p = doc.add_heading(titulo_alerta, 3)

                # Color según nivel
                if p.runs:
                    color_map = {
                        'CRITICA': RGBColor(204, 51, 51),
                        'ALTA': RGBColor(230, 115, 0),
                        'MEDIA': RGBColor(204, 204, 0),
                        'BAJA': RGBColor(51, 204, 51),
                    }
                    if nivel in color_map:
                        p.runs[0].font.color.rgb = color_map[nivel]

                doc.add_paragraph(alerta.get('descripcion', ''))

        # Normativa citada
        normativa = informe.get('normativa_citada', [])
        if normativa:
            doc.add_heading('NORMATIVA CITADA', 2)
            for norm in normativa[:15]:
                ref = norm.get('referencia', '') or norm.get('documento', '')
                if ref:
                    doc.add_paragraph(ref, style='List Bullet')

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "Este informe es un análisis preliminar de prefactibilidad ambiental. "
            "No reemplaza la evaluación formal del SEA ni la opinión de expertos especializados."
        ).italic = True

        # Guardar a bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _exportar_txt(self, informe: dict[str, Any]) -> bytes:
        """Genera el informe en formato texto plano."""
        lineas = []

        lineas.append("=" * 80)
        lineas.append("INFORME DE PREFACTIBILIDAD AMBIENTAL")
        lineas.append("=" * 80)

        datos_proyecto = informe.get('datos_proyecto', {})
        lineas.append(f"\nProyecto: {datos_proyecto.get('nombre', 'N/D')}")
        lineas.append(f"ID: {informe.get('id', 'N/D')}")
        lineas.append(f"Fecha: {informe.get('fecha_generacion', 'N/D')}")

        clasif = informe.get('clasificacion_seia', {})
        lineas.append(f"\nVía Recomendada: {clasif.get('via_ingreso_recomendada', 'N/D')}")
        lineas.append(f"Confianza: {clasif.get('confianza', 0) * 100:.0f}%")

        lineas.append("\n" + "-" * 80)

        # Secciones
        secciones = informe.get('secciones', [])
        for seccion in secciones:
            if isinstance(seccion, dict):
                lineas.append(f"\n## {seccion.get('titulo', 'Sin Título').upper()}\n")
                lineas.append(seccion.get('contenido', ''))
                lineas.append("\n" + "-" * 40)

        # Alertas
        alertas = informe.get('alertas', [])
        if alertas:
            lineas.append("\n## ALERTAS AMBIENTALES\n")
            for alerta in alertas[:10]:
                lineas.append(f"[{alerta.get('nivel', 'INFO')}] {alerta.get('titulo', 'N/D')}")
                lineas.append(f"  {alerta.get('descripcion', '')}")
                lineas.append("")

        lineas.append("\n" + "=" * 80)
        lineas.append("FIN DEL INFORME")
        lineas.append("=" * 80)

        texto = "\n".join(lineas)
        return texto.encode('utf-8')

    def _exportar_html(self, informe: dict[str, Any]) -> bytes:
        """Genera el informe en formato HTML."""
        datos_proyecto = informe.get('datos_proyecto', {})
        clasif = informe.get('clasificacion_seia', {})

        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="es">',
            '<head>',
            '  <meta charset="UTF-8">',
            '  <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f'  <title>Informe Prefactibilidad - {datos_proyecto.get("nombre", "Proyecto")}</title>',
            '  <style>',
            '    body { font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }',
            '    h1 { color: #1a4d7c; text-align: center; }',
            '    h2 { color: #2d6da3; border-bottom: 2px solid #2d6da3; padding-bottom: 5px; }',
            '    h3 { color: #3d7db3; }',
            '    .info-table { border-collapse: collapse; margin: 20px 0; }',
            '    .info-table td { padding: 8px 15px; border: 1px solid #ddd; }',
            '    .info-table td:first-child { font-weight: bold; background: #f5f5f5; }',
            '    .alerta { padding: 10px; margin: 10px 0; border-left: 4px solid; }',
            '    .alerta-CRITICA { border-color: #cc3333; background: #ffeeee; }',
            '    .alerta-ALTA { border-color: #e67300; background: #fff5ee; }',
            '    .alerta-MEDIA { border-color: #cccc00; background: #ffffee; }',
            '    .alerta-BAJA { border-color: #33cc33; background: #eeffee; }',
            '    .disclaimer { font-style: italic; color: #666; margin-top: 30px; padding: 15px; background: #f9f9f9; }',
            '    .seccion { margin: 20px 0; }',
            '  </style>',
            '</head>',
            '<body>',
            '  <h1>INFORME DE PREFACTIBILIDAD AMBIENTAL</h1>',
            f'  <h2 style="text-align: center;">{datos_proyecto.get("nombre", "Proyecto Sin Nombre")}</h2>',
            '  <table class="info-table">',
            f'    <tr><td>ID Informe</td><td>{informe.get("id", "N/D")}</td></tr>',
            f'    <tr><td>Fecha</td><td>{informe.get("fecha_generacion", "N/D")}</td></tr>',
            f'    <tr><td>Vía Recomendada</td><td><strong>{clasif.get("via_ingreso_recomendada", "N/D")}</strong></td></tr>',
            f'    <tr><td>Confianza</td><td>{clasif.get("confianza", 0) * 100:.0f}%</td></tr>',
            '  </table>',
        ]

        # Secciones
        secciones = informe.get('secciones', [])
        for seccion in secciones:
            if isinstance(seccion, dict):
                html_parts.append(f'  <div class="seccion">')
                html_parts.append(f'    <h2>{seccion.get("titulo", "Sin Título")}</h2>')
                contenido = seccion.get('contenido', '').replace('\n', '<br>')
                html_parts.append(f'    <p>{contenido}</p>')
                html_parts.append('  </div>')

        # Alertas
        alertas = informe.get('alertas', [])
        if alertas:
            html_parts.append('  <h2>ALERTAS AMBIENTALES</h2>')
            for alerta in alertas[:10]:
                nivel = alerta.get('nivel', 'INFO')
                html_parts.append(f'  <div class="alerta alerta-{nivel}">')
                html_parts.append(f'    <strong>[{nivel}] {alerta.get("titulo", "N/D")}</strong>')
                html_parts.append(f'    <p>{alerta.get("descripcion", "")}</p>')
                html_parts.append('  </div>')

        # Normativa
        normativa = informe.get('normativa_citada', [])
        if normativa:
            html_parts.append('  <h2>NORMATIVA CITADA</h2>')
            html_parts.append('  <ul>')
            for norm in normativa[:15]:
                ref = norm.get('referencia', '') or norm.get('documento', '')
                if ref:
                    html_parts.append(f'    <li>{ref}</li>')
            html_parts.append('  </ul>')

        html_parts.extend([
            '  <div class="disclaimer">',
            '    Este informe es un análisis preliminar de prefactibilidad ambiental. ',
            '    No reemplaza la evaluación formal del SEA ni la opinión de expertos especializados.',
            '  </div>',
            '</body>',
            '</html>',
        ])

        html = '\n'.join(html_parts)
        return html.encode('utf-8')

    def obtener_formatos_disponibles(self) -> list[dict]:
        """Retorna los formatos de exportación disponibles."""
        return [
            {
                "formato": FormatoExportacion.PDF.value,
                "extension": ".pdf",
                "mime_type": "application/pdf",
                "descripcion": "Documento PDF formal",
            },
            {
                "formato": FormatoExportacion.DOCX.value,
                "extension": ".docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "descripcion": "Documento Microsoft Word editable",
            },
            {
                "formato": FormatoExportacion.TXT.value,
                "extension": ".txt",
                "mime_type": "text/plain",
                "descripcion": "Texto plano",
            },
            {
                "formato": FormatoExportacion.HTML.value,
                "extension": ".html",
                "mime_type": "text/html",
                "descripcion": "Documento HTML",
            },
        ]
