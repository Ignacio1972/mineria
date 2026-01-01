"""
Servicio para gestion de archivos originales del corpus RAG.

Responsabilidades:
- Almacenar archivos con estructura año/mes/uuid
- Calcular hash SHA-256 para integridad y deduplicacion
- Extraer texto de PDFs y DOCX
- Gestionar eliminacion segura
"""

import hashlib
import logging
from pathlib import Path
from datetime import datetime
from uuid import uuid4
from typing import Optional, Tuple, BinaryIO

import aiofiles
import aiofiles.os

from app.core.config import settings

logger = logging.getLogger(__name__)


class ArchivoService:
    """
    Servicio para gestionar archivos originales (PDF/DOCX) del corpus RAG.
    """

    # Tipos MIME permitidos
    MIME_TYPES_PERMITIDOS = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    }

    def __init__(self, base_path: Optional[str] = None):
        """
        Inicializa el servicio de archivos.

        Args:
            base_path: Ruta base para almacenamiento. Por defecto usa settings.
        """
        self.base_path = Path(base_path or getattr(settings, 'STORAGE_PATH', '/var/www/mineria/storage/legal'))
        self.base_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"ArchivoService inicializado en: {self.base_path}")

    def validar_tipo_archivo(self, mime_type: str) -> bool:
        """Valida si el tipo MIME es permitido."""
        return mime_type in self.MIME_TYPES_PERMITIDOS

    def calcular_hash(self, contenido: bytes) -> str:
        """Calcula el hash SHA-256 del contenido."""
        return hashlib.sha256(contenido).hexdigest()

    def generar_ruta_storage(self, extension: str) -> Tuple[str, str]:
        """
        Genera una ruta de almacenamiento unica.

        Returns:
            Tuple con (ruta_relativa, nombre_storage)
        """
        ahora = datetime.now()
        nombre_storage = f"{uuid4()}{extension}"
        ruta_relativa = f"{ahora.year}/{ahora.month:02d}/{nombre_storage}"
        return ruta_relativa, nombre_storage

    async def guardar_archivo(
        self,
        contenido: bytes,
        nombre_original: str,
        mime_type: str,
        subido_por: Optional[str] = None
    ) -> Tuple[str, str, int, str]:
        """
        Guarda un archivo en el storage.

        Args:
            contenido: Bytes del archivo
            nombre_original: Nombre original del archivo
            mime_type: Tipo MIME del archivo
            subido_por: Usuario que sube el archivo

        Returns:
            Tuple con (ruta_storage, hash_sha256, tamano_bytes, nombre_storage)

        Raises:
            ValueError: Si el tipo de archivo no es permitido
        """
        # Validar tipo
        if not self.validar_tipo_archivo(mime_type):
            raise ValueError(f"Tipo de archivo no permitido: {mime_type}")

        # Calcular hash
        hash_sha256 = self.calcular_hash(contenido)

        # Obtener extension
        extension = self.MIME_TYPES_PERMITIDOS[mime_type]

        # Generar ruta
        ruta_relativa, nombre_storage = self.generar_ruta_storage(extension)
        ruta_completa = self.base_path / ruta_relativa

        # Crear directorio si no existe
        await aiofiles.os.makedirs(ruta_completa.parent, exist_ok=True)

        # Guardar archivo
        async with aiofiles.open(ruta_completa, 'wb') as f:
            await f.write(contenido)

        logger.info(f"Archivo guardado: {ruta_relativa} ({len(contenido)} bytes)")

        return ruta_relativa, hash_sha256, len(contenido), nombre_storage

    async def guardar_desde_upload(
        self,
        archivo: BinaryIO,
        nombre_original: str,
        mime_type: str,
        subido_por: Optional[str] = None
    ) -> Tuple[str, str, int, str]:
        """
        Guarda un archivo desde un objeto UploadFile de FastAPI.

        Args:
            archivo: Objeto archivo (UploadFile.file)
            nombre_original: Nombre original del archivo
            mime_type: Tipo MIME
            subido_por: Usuario que sube

        Returns:
            Tuple con (ruta_storage, hash_sha256, tamano_bytes, nombre_storage)
        """
        contenido = archivo.read()
        return await self.guardar_archivo(contenido, nombre_original, mime_type, subido_por)

    async def obtener_ruta_completa(self, ruta_storage: str) -> Path:
        """Retorna la ruta absoluta de un archivo."""
        return self.base_path / ruta_storage

    async def archivo_existe(self, ruta_storage: str) -> bool:
        """Verifica si un archivo existe."""
        ruta = self.base_path / ruta_storage
        return ruta.exists()

    async def eliminar_archivo(self, ruta_storage: str) -> bool:
        """
        Elimina un archivo del storage.

        Returns:
            True si se elimino, False si no existia
        """
        ruta = self.base_path / ruta_storage
        if ruta.exists():
            await aiofiles.os.remove(ruta)
            logger.info(f"Archivo eliminado: {ruta_storage}")
            return True
        return False

    async def leer_archivo(self, ruta_storage: str) -> bytes:
        """Lee y retorna el contenido de un archivo."""
        ruta = self.base_path / ruta_storage
        if not ruta.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {ruta_storage}")

        async with aiofiles.open(ruta, 'rb') as f:
            return await f.read()

    def extraer_texto_pdf(self, ruta_storage: str) -> Tuple[str, int]:
        """
        Extrae texto de un PDF usando PyMuPDF.

        Args:
            ruta_storage: Ruta relativa del archivo

        Returns:
            Tuple con (texto_extraido, numero_paginas)
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF no instalado. Instalar con: pip install PyMuPDF")
            raise ImportError("PyMuPDF no esta instalado")

        ruta = self.base_path / ruta_storage
        texto_completo = []

        with fitz.open(str(ruta)) as doc:
            num_paginas = len(doc)
            for pagina in doc:
                texto = pagina.get_text()
                if texto.strip():
                    texto_completo.append(texto)

        texto_final = "\n\n".join(texto_completo)
        logger.info(f"Texto extraido de PDF: {len(texto_final)} caracteres, {num_paginas} paginas")

        return texto_final, num_paginas

    def extraer_texto_docx(self, ruta_storage: str) -> str:
        """
        Extrae texto de un archivo DOCX.

        Args:
            ruta_storage: Ruta relativa del archivo

        Returns:
            Texto extraido del documento
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx no instalado. Instalar con: pip install python-docx")
            raise ImportError("python-docx no esta instalado")

        ruta = self.base_path / ruta_storage
        doc = Document(str(ruta))

        texto_completo = []
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                texto_completo.append(parrafo.text)

        # Tambien extraer texto de tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text.strip():
                        texto_completo.append(celda.text)

        texto_final = "\n\n".join(texto_completo)
        logger.info(f"Texto extraido de DOCX: {len(texto_final)} caracteres")

        return texto_final

    def extraer_texto(self, ruta_storage: str, mime_type: str) -> Tuple[str, Optional[int]]:
        """
        Extrae texto de un archivo segun su tipo.

        Args:
            ruta_storage: Ruta relativa del archivo
            mime_type: Tipo MIME del archivo

        Returns:
            Tuple con (texto_extraido, numero_paginas o None)
        """
        if mime_type == "application/pdf":
            return self.extraer_texto_pdf(ruta_storage)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            texto = self.extraer_texto_docx(ruta_storage)
            return texto, None
        else:
            raise ValueError(f"Tipo de archivo no soportado para extraccion: {mime_type}")

    async def obtener_estadisticas(self) -> dict:
        """
        Obtiene estadisticas del storage.

        Returns:
            Dict con estadisticas (total_archivos, tamano_total, etc.)
        """
        import os

        total_archivos = 0
        tamano_total = 0
        por_extension = {}

        for root, dirs, files in os.walk(self.base_path):
            for file in files:
                total_archivos += 1
                file_path = Path(root) / file
                tamano = file_path.stat().st_size
                tamano_total += tamano

                ext = file_path.suffix.lower()
                por_extension[ext] = por_extension.get(ext, 0) + 1

        return {
            "total_archivos": total_archivos,
            "tamano_total_bytes": tamano_total,
            "tamano_total_mb": round(tamano_total / (1024 * 1024), 2),
            "por_extension": por_extension,
            "ruta_base": str(self.base_path),
        }


# Instancia singleton del servicio
_archivo_service: Optional[ArchivoService] = None


def get_archivo_service() -> ArchivoService:
    """
    Obtiene la instancia singleton del servicio de archivos.

    Para usar como dependencia de FastAPI:
        archivo_service: ArchivoService = Depends(get_archivo_service)
    """
    global _archivo_service
    if _archivo_service is None:
        _archivo_service = ArchivoService()
    return _archivo_service
