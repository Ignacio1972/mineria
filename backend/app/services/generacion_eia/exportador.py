"""
Servicio de Exportación de Documentos EIA.

Genera documentos en formatos PDF, DOCX y e-SEIA XML.
"""
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
from io import BytesIO

import markdown
from jinja2 import Environment, FileSystemLoader
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.config import settings
from app.db.models.generacion_eia import (
    DocumentoEIA,
    ExportacionEIA,
    FormatoExportacion
)
from app.db.models import Proyecto
from app.schemas.generacion_eia import (
    ConfiguracionPDF,
    ConfiguracionDOCX,
    ConfiguracionESEIA,
    ExportacionResponse
)

logger = logging.getLogger(__name__)

# Directorio base para exportaciones
EXPORTS_DIR = Path("/app/exports")
TEMPLATES_DIR = Path(__file__).parent / "templates"


class ExportadorService:
    """Servicio para exportar documentos EIA a diferentes formatos."""

    def __init__(self):
        """Inicializa el servicio de exportación."""
        self._jinja_env: Optional[Environment] = None
        self._md = markdown.Markdown(
            extensions=[
                'tables',
                'fenced_code',
                'toc',
                'attr_list',
                'def_list',
                'footnotes'
            ]
        )
        # Asegurar que existe el directorio de exports
        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

    @property
    def jinja_env(self) -> Environment:
        """Entorno Jinja2 lazy-loaded."""
        if self._jinja_env is None:
            self._jinja_env = Environment(
                loader=FileSystemLoader(str(TEMPLATES_DIR)),
                autoescape=True
            )
        return self._jinja_env

    # =========================================================================
    # MÉTODOS PÚBLICOS
    # =========================================================================

    async def exportar_pdf(
        self,
        db: AsyncSession,
        documento_id: int,
        config: Optional[ConfiguracionPDF] = None
    ) -> ExportacionResponse:
        """
        Exporta un documento EIA a formato PDF.

        Args:
            db: Sesión de base de datos
            documento_id: ID del documento EIA
            config: Configuración de exportación PDF

        Returns:
            ExportacionResponse con información del archivo generado
        """
        logger.info(f"Exportando documento {documento_id} a PDF")

        if config is None:
            config = ConfiguracionPDF()

        try:
            # 1. Obtener documento y proyecto
            documento = await self._get_documento(db, documento_id)
            proyecto = await self._get_proyecto(db, documento.proyecto_id)

            # 2. Preparar datos para template
            datos_template = await self._preparar_datos_template(
                documento, proyecto, config
            )

            # 3. Renderizar HTML
            html_content = self._renderizar_html(datos_template, config)

            # 4. Generar PDF con WeasyPrint
            pdf_bytes = self._generar_pdf_weasyprint(html_content)

            # 5. Guardar archivo
            filename = self._generar_nombre_archivo(proyecto, documento, "pdf")
            filepath = EXPORTS_DIR / filename
            filepath.write_bytes(pdf_bytes)

            # 6. Registrar exportación
            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.PDF,
                filepath=str(filepath),
                tamano=len(pdf_bytes),
                config=config.model_dump()
            )

            logger.info(f"PDF generado exitosamente: {filepath} ({len(pdf_bytes)} bytes)")

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.PDF,
                archivo_path=str(filepath),
                tamano_bytes=len(pdf_bytes),
                generado_exitoso=True,
                error_mensaje=None,
                created_at=exportacion.created_at,
                url_descarga=f"/api/v1/generacion/{documento.proyecto_id}/export/{exportacion.id}"
            )

        except Exception as e:
            logger.error(f"Error exportando PDF: {e}", exc_info=True)

            # Registrar fallo
            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.PDF,
                filepath=None,
                tamano=0,
                config=config.model_dump() if config else {},
                error=str(e)
            )

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.PDF,
                archivo_path=None,
                tamano_bytes=None,
                generado_exitoso=False,
                error_mensaje=str(e),
                created_at=exportacion.created_at,
                url_descarga=None
            )

    async def exportar_docx(
        self,
        db: AsyncSession,
        documento_id: int,
        config: Optional[ConfiguracionDOCX] = None
    ) -> ExportacionResponse:
        """
        Exporta un documento EIA a formato DOCX.

        Args:
            db: Sesión de base de datos
            documento_id: ID del documento EIA
            config: Configuración de exportación DOCX

        Returns:
            ExportacionResponse con información del archivo generado
        """
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        logger.info(f"Exportando documento {documento_id} a DOCX")

        if config is None:
            config = ConfiguracionDOCX()

        try:
            # 1. Obtener documento y proyecto
            documento = await self._get_documento(db, documento_id)
            proyecto = await self._get_proyecto(db, documento.proyecto_id)

            # 2. Crear documento Word
            doc = Document()

            # 3. Configurar estilos
            self._configurar_estilos_docx(doc)

            # 4. Agregar portada si se requiere
            if config.incluir_portada:
                self._agregar_portada_docx(doc, proyecto, documento)
                doc.add_page_break()

            # 5. Agregar índice si se requiere
            if config.incluir_indice:
                self._agregar_indice_docx(doc, documento)
                doc.add_page_break()

            # 6. Agregar capítulos
            capitulos = documento.contenido_capitulos or {}
            for num_str, capitulo_data in sorted(capitulos.items(), key=lambda x: int(x[0])):
                self._agregar_capitulo_docx(doc, int(num_str), capitulo_data)

            # 7. Guardar documento
            filename = self._generar_nombre_archivo(proyecto, documento, "docx")
            filepath = EXPORTS_DIR / filename

            doc.save(str(filepath))
            tamano = filepath.stat().st_size

            # 8. Registrar exportación
            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.DOCX,
                filepath=str(filepath),
                tamano=tamano,
                config=config.model_dump()
            )

            logger.info(f"DOCX generado exitosamente: {filepath} ({tamano} bytes)")

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.DOCX,
                archivo_path=str(filepath),
                tamano_bytes=tamano,
                generado_exitoso=True,
                error_mensaje=None,
                created_at=exportacion.created_at,
                url_descarga=f"/api/v1/generacion/{documento.proyecto_id}/export/{exportacion.id}"
            )

        except Exception as e:
            logger.error(f"Error exportando DOCX: {e}", exc_info=True)

            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.DOCX,
                filepath=None,
                tamano=0,
                config=config.model_dump() if config else {},
                error=str(e)
            )

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.DOCX,
                archivo_path=None,
                tamano_bytes=None,
                generado_exitoso=False,
                error_mensaje=str(e),
                created_at=exportacion.created_at,
                url_descarga=None
            )

    async def exportar_eseia(
        self,
        db: AsyncSession,
        documento_id: int,
        config: Optional[ConfiguracionESEIA] = None
    ) -> ExportacionResponse:
        """
        Exporta un documento EIA a formato e-SEIA XML.

        Args:
            db: Sesión de base de datos
            documento_id: ID del documento EIA
            config: Configuración de exportación e-SEIA

        Returns:
            ExportacionResponse con información del archivo generado
        """
        from lxml import etree

        logger.info(f"Exportando documento {documento_id} a e-SEIA XML")

        if config is None:
            config = ConfiguracionESEIA()

        try:
            # 1. Obtener documento y proyecto
            documento = await self._get_documento(db, documento_id)
            proyecto = await self._get_proyecto(db, documento.proyecto_id)

            # 2. Construir estructura XML
            root = self._construir_xml_eseia(documento, proyecto, config)

            # 3. Serializar XML
            xml_bytes = etree.tostring(
                root,
                encoding='utf-8',
                xml_declaration=True,
                pretty_print=True
            )

            # 4. Guardar archivo
            filename = self._generar_nombre_archivo(proyecto, documento, "xml")
            filepath = EXPORTS_DIR / filename
            filepath.write_bytes(xml_bytes)

            # 5. Registrar exportación
            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.ESEIA_XML,
                filepath=str(filepath),
                tamano=len(xml_bytes),
                config=config.model_dump()
            )

            logger.info(f"XML e-SEIA generado exitosamente: {filepath}")

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.ESEIA_XML,
                archivo_path=str(filepath),
                tamano_bytes=len(xml_bytes),
                generado_exitoso=True,
                error_mensaje=None,
                created_at=exportacion.created_at,
                url_descarga=f"/api/v1/generacion/{documento.proyecto_id}/export/{exportacion.id}"
            )

        except Exception as e:
            logger.error(f"Error exportando e-SEIA XML: {e}", exc_info=True)

            exportacion = await self._registrar_exportacion(
                db=db,
                documento_id=documento_id,
                formato=FormatoExportacion.ESEIA_XML,
                filepath=None,
                tamano=0,
                config=config.model_dump() if config else {},
                error=str(e)
            )

            return ExportacionResponse(
                id=exportacion.id,
                documento_id=documento_id,
                formato=FormatoExportacion.ESEIA_XML,
                archivo_path=None,
                tamano_bytes=None,
                generado_exitoso=False,
                error_mensaje=str(e),
                created_at=exportacion.created_at,
                url_descarga=None
            )

    async def generar_indice(
        self,
        db: AsyncSession,
        documento_id: int
    ) -> List[Dict[str, Any]]:
        """
        Genera el índice de contenidos de un documento EIA.

        Args:
            db: Sesión de base de datos
            documento_id: ID del documento EIA

        Returns:
            Lista de items del índice con nivel, número, título y página estimada
        """
        documento = await self._get_documento(db, documento_id)
        capitulos = documento.contenido_capitulos or {}

        indice = []
        pagina_estimada = 3  # Portada + índice

        for num_str, capitulo_data in sorted(capitulos.items(), key=lambda x: int(x[0])):
            num = int(num_str)
            titulo = capitulo_data.get('titulo', f'Capítulo {num}')
            contenido = capitulo_data.get('contenido', '')

            # Agregar capítulo principal
            indice.append({
                'nivel': 1,
                'numero': f'{num}.',
                'titulo': titulo,
                'pagina': pagina_estimada
            })

            # Buscar subsecciones (## en markdown)
            subsecciones = capitulo_data.get('subsecciones', {})
            seccion_num = 1
            for seccion_titulo in subsecciones.keys():
                indice.append({
                    'nivel': 2,
                    'numero': f'{num}.{seccion_num}',
                    'titulo': seccion_titulo,
                    'pagina': pagina_estimada + 1
                })
                seccion_num += 1

            # Estimar páginas basado en contenido
            palabras = len(contenido.split()) if contenido else 0
            paginas_capitulo = max(1, palabras // 400)  # ~400 palabras por página
            pagina_estimada += paginas_capitulo

        return indice

    async def get_exportacion(
        self,
        db: AsyncSession,
        exportacion_id: int
    ) -> Optional[ExportacionEIA]:
        """
        Obtiene una exportación por su ID.

        Args:
            db: Sesión de base de datos
            exportacion_id: ID de la exportación

        Returns:
            ExportacionEIA o None
        """
        query = select(ExportacionEIA).where(ExportacionEIA.id == exportacion_id)
        result = await db.execute(query)
        return result.scalar_one_or_none()

    async def listar_exportaciones(
        self,
        db: AsyncSession,
        documento_id: int
    ) -> List[ExportacionEIA]:
        """
        Lista todas las exportaciones de un documento.

        Args:
            db: Sesión de base de datos
            documento_id: ID del documento EIA

        Returns:
            Lista de ExportacionEIA
        """
        query = select(ExportacionEIA).where(
            ExportacionEIA.documento_id == documento_id
        ).order_by(ExportacionEIA.created_at.desc())

        result = await db.execute(query)
        return list(result.scalars().all())

    async def get_archivo_exportacion(
        self,
        db: AsyncSession,
        exportacion_id: int
    ) -> Optional[tuple[bytes, str, str]]:
        """
        Obtiene el contenido de un archivo exportado.

        Args:
            db: Sesión de base de datos
            exportacion_id: ID de la exportación

        Returns:
            Tuple de (bytes, filename, content_type) o None
        """
        exportacion = await self.get_exportacion(db, exportacion_id)

        if not exportacion or not exportacion.archivo_path:
            return None

        filepath = Path(exportacion.archivo_path)
        if not filepath.exists():
            logger.error(f"Archivo no encontrado: {filepath}")
            return None

        content_types = {
            FormatoExportacion.PDF.value: "application/pdf",
            FormatoExportacion.DOCX.value: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            FormatoExportacion.ESEIA_XML.value: "application/xml"
        }

        content_type = content_types.get(exportacion.formato, "application/octet-stream")
        filename = filepath.name

        return filepath.read_bytes(), filename, content_type

    # =========================================================================
    # MÉTODOS PRIVADOS - PDF
    # =========================================================================

    async def _preparar_datos_template(
        self,
        documento: DocumentoEIA,
        proyecto: Proyecto,
        config: ConfiguracionPDF
    ) -> Dict[str, Any]:
        """Prepara los datos para el template HTML."""
        capitulos = documento.contenido_capitulos or {}
        metadatos = documento.metadatos or {}

        # Convertir contenido markdown a HTML por capítulo
        capitulos_html = []
        for num_str, capitulo_data in sorted(capitulos.items(), key=lambda x: int(x[0])):
            contenido_md = capitulo_data.get('contenido', '')
            self._md.reset()
            contenido_html = self._md.convert(contenido_md)

            capitulos_html.append({
                'numero': int(num_str),
                'titulo': capitulo_data.get('titulo', f'Capítulo {num_str}'),
                'contenido_html': contenido_html
            })

        # Generar índice
        indice = []
        for cap in capitulos_html:
            indice.append({
                'nivel': 1,
                'numero': f"{cap['numero']}.",
                'titulo': cap['titulo'],
                'pagina': ''  # WeasyPrint no soporta número de página dinámico fácilmente
            })

        return {
            'titulo': documento.titulo or f"EIA - {proyecto.nombre}",
            'titulo_corto': proyecto.nombre[:50] if proyecto.nombre else 'EIA',
            'nombre_proyecto': proyecto.nombre,
            'titular': proyecto.titular or 'N/A',
            'region': proyecto.region or 'N/A',
            'comuna': proyecto.comuna or 'N/A',
            'tipo_proyecto': proyecto.descripcion or 'N/A',
            'empresa_consultora': metadatos.get('empresa_consultora', ''),
            'fecha_elaboracion': metadatos.get('fecha_elaboracion', datetime.now().strftime('%d/%m/%Y')),
            'version': f"v{documento.version}",
            'capitulos': capitulos_html,
            'indice': indice,
            'anexos': [],  # TODO: Implementar anexos
            # Configuración
            'incluir_portada': config.incluir_portada,
            'incluir_indice': config.incluir_indice,
            'tamano_pagina': config.tamano_pagina,
            'orientacion': config.orientacion,
            'margen_cm': config.margen_cm,
            'fuente': config.fuente,
            'tamano_fuente': config.tamano_fuente,
            'logo_empresa': config.logo_empresa
        }

    def _renderizar_html(
        self,
        datos: Dict[str, Any],
        config: ConfiguracionPDF
    ) -> str:
        """Renderiza el template HTML con los datos."""
        template = self.jinja_env.get_template('eia_pdf.html')
        return template.render(**datos)

    def _generar_pdf_weasyprint(self, html_content: str) -> bytes:
        """
        Genera PDF usando WeasyPrint si está disponible,
        de lo contrario usa reportlab como fallback.
        """
        try:
            from weasyprint import HTML
            html_doc = HTML(string=html_content, base_url=str(TEMPLATES_DIR))
            return html_doc.write_pdf()
        except (ImportError, OSError) as e:
            logger.warning(f"WeasyPrint no disponible ({e}), usando reportlab")
            return self._generar_pdf_reportlab(html_content)

    def _generar_pdf_reportlab(self, html_content: str) -> bytes:
        """Genera PDF usando reportlab (fallback)."""
        from io import BytesIO
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak,
            Table, TableStyle, Image
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.colors import HexColor
        import re

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2.5*cm,
            leftMargin=2.5*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )

        # Estilos personalizados
        styles = getSampleStyleSheet()

        # Título principal (portada)
        styles.add(ParagraphStyle(
            name='CoverTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=HexColor('#1a365d')
        ))

        # Subtítulo portada
        styles.add(ParagraphStyle(
            name='CoverSubtitle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=40,
            alignment=TA_CENTER,
            textColor=HexColor('#2c5282')
        ))

        # Nombre del proyecto
        styles.add(ParagraphStyle(
            name='ProjectName',
            parent=styles['Normal'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=40,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            textColor=HexColor('#1a365d')
        ))

        # Título de capítulo
        styles.add(ParagraphStyle(
            name='ChapterTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=15,
            textColor=HexColor('#1a365d'),
            borderWidth=1,
            borderColor=HexColor('#2c5282'),
            borderPadding=5
        ))

        # Subtítulo (H2)
        styles.add(ParagraphStyle(
            name='SectionTitle',
            parent=styles['Heading2'],
            fontSize=13,
            spaceBefore=15,
            spaceAfter=10,
            textColor=HexColor('#2c5282')
        ))

        # Texto normal con justificación
        styles.add(ParagraphStyle(
            name='BodyJustified',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceBefore=6,
            spaceAfter=6
        ))

        # Lista
        styles.add(ParagraphStyle(
            name='ListItem',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceBefore=3,
            spaceAfter=3
        ))

        # Información de portada
        styles.add(ParagraphStyle(
            name='CoverInfo',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceBefore=8,
            spaceAfter=8
        ))

        story = []

        # Extraer información del HTML
        info = self._extraer_info_html(html_content)

        # === PORTADA ===
        if info.get('incluir_portada', True):
            story.append(Spacer(1, 3*cm))
            story.append(Paragraph("ESTUDIO DE IMPACTO AMBIENTAL", styles['CoverTitle']))
            story.append(Paragraph("Conforme a la Ley 19.300 y D.S. 40/2012", styles['CoverSubtitle']))
            story.append(Spacer(1, 1*cm))
            story.append(Paragraph(info.get('nombre_proyecto', 'Proyecto'), styles['ProjectName']))
            story.append(Spacer(1, 2*cm))

            # Info del proyecto
            if info.get('titular'):
                story.append(Paragraph(f"<b>Titular:</b> {info['titular']}", styles['CoverInfo']))
            if info.get('ubicacion'):
                story.append(Paragraph(f"<b>Ubicación:</b> {info['ubicacion']}", styles['CoverInfo']))
            if info.get('empresa_consultora'):
                story.append(Paragraph(f"<b>Consultora:</b> {info['empresa_consultora']}", styles['CoverInfo']))

            story.append(Spacer(1, 2*cm))

            if info.get('fecha'):
                story.append(Paragraph(f"<b>Fecha:</b> {info['fecha']}", styles['CoverInfo']))
            if info.get('version'):
                story.append(Paragraph(f"<b>Versión:</b> {info['version']}", styles['CoverInfo']))

            story.append(PageBreak())

        # === ÍNDICE ===
        if info.get('incluir_indice', True) and info.get('capitulos'):
            story.append(Paragraph("ÍNDICE DE CONTENIDOS", styles['ChapterTitle']))
            story.append(Spacer(1, 0.5*cm))

            for cap in info.get('capitulos', []):
                story.append(Paragraph(
                    f"<b>{cap.get('numero', '')}.</b> {cap.get('titulo', '')}",
                    styles['BodyJustified']
                ))

            story.append(PageBreak())

        # === CAPÍTULOS ===
        for cap in info.get('capitulos', []):
            story.append(Paragraph(
                f"Capítulo {cap.get('numero', '')}: {cap.get('titulo', '')}",
                styles['ChapterTitle']
            ))

            contenido = cap.get('contenido_html', '')
            paragraphs = self._html_a_paragraphs(contenido, styles)
            story.extend(paragraphs)

            story.append(PageBreak())

        # Generar PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()

        return pdf_bytes

    def _extraer_info_html(self, html: str) -> Dict[str, Any]:
        """Extrae información estructurada del HTML."""
        import re

        info = {
            'incluir_portada': True,
            'incluir_indice': True,
            'nombre_proyecto': '',
            'titular': '',
            'ubicacion': '',
            'empresa_consultora': '',
            'fecha': '',
            'version': '',
            'capitulos': []
        }

        # Extraer nombre del proyecto
        match = re.search(r'<div class="cover-project">([^<]+)</div>', html)
        if match:
            info['nombre_proyecto'] = match.group(1).strip()

        # Extraer titular
        match = re.search(r'<strong>Titular:</strong>\s*([^<]+)', html)
        if match:
            info['titular'] = match.group(1).strip()

        # Extraer ubicación
        match = re.search(r'<strong>Ubicación:</strong>\s*([^<]+)', html)
        if match:
            info['ubicacion'] = match.group(1).strip()

        # Extraer consultora
        match = re.search(r'<strong>Consultora:</strong>\s*([^<]+)', html)
        if match:
            info['empresa_consultora'] = match.group(1).strip()

        # Extraer fecha
        match = re.search(r'<strong>Fecha:</strong>\s*([^<]+)', html)
        if match:
            info['fecha'] = match.group(1).strip()

        # Extraer versión
        match = re.search(r'<strong>Versión:</strong>\s*([^<]+)', html)
        if match:
            info['version'] = match.group(1).strip()

        # Extraer capítulos
        chapter_pattern = re.compile(
            r'<div class="chapter"[^>]*>.*?'
            r'<div class="chapter-number">Capítulo (\d+)</div>\s*'
            r'<h1>([^<]+)</h1>\s*'
            r'(.*?)</div>',
            re.DOTALL
        )

        for match in chapter_pattern.finditer(html):
            info['capitulos'].append({
                'numero': int(match.group(1)),
                'titulo': match.group(2).strip(),
                'contenido_html': match.group(3).strip()
            })

        return info

    def _html_a_paragraphs(self, html: str, styles) -> list:
        """Convierte contenido HTML a párrafos de reportlab."""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import cm
        import re

        paragraphs = []

        # Limpiar HTML
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)

        # Procesar párrafos
        for p_match in re.finditer(r'<p>(.+?)</p>', html, re.DOTALL):
            texto = p_match.group(1).strip()
            texto = self._limpiar_html_basico(texto)
            if texto:
                paragraphs.append(Paragraph(texto, styles['BodyJustified']))

        # Procesar headers h2
        for h_match in re.finditer(r'<h2>(.+?)</h2>', html, re.DOTALL):
            texto = h_match.group(1).strip()
            texto = self._limpiar_html_basico(texto)
            if texto:
                paragraphs.append(Spacer(1, 0.3*cm))
                paragraphs.append(Paragraph(texto, styles['SectionTitle']))

        # Procesar headers h3
        for h_match in re.finditer(r'<h3>(.+?)</h3>', html, re.DOTALL):
            texto = h_match.group(1).strip()
            texto = self._limpiar_html_basico(texto)
            if texto:
                paragraphs.append(Paragraph(f"<b>{texto}</b>", styles['BodyJustified']))

        # Procesar listas
        for li_match in re.finditer(r'<li>(.+?)</li>', html, re.DOTALL):
            texto = li_match.group(1).strip()
            texto = self._limpiar_html_basico(texto)
            if texto:
                paragraphs.append(Paragraph(f"• {texto}", styles['ListItem']))

        if not paragraphs:
            # Si no hay contenido estructurado, procesar como texto plano
            texto = re.sub(r'<[^>]+>', ' ', html)
            texto = ' '.join(texto.split())
            if texto:
                paragraphs.append(Paragraph(texto, styles['BodyJustified']))

        return paragraphs

    def _limpiar_html_basico(self, texto: str) -> str:
        """Limpia HTML básico preservando negritas y cursivas."""
        import re
        # Convertir <strong> y <b> a negritas de reportlab
        texto = re.sub(r'<strong>(.+?)</strong>', r'<b>\1</b>', texto)
        # Convertir <em> y <i> a cursivas de reportlab
        texto = re.sub(r'<em>(.+?)</em>', r'<i>\1</i>', texto)
        # Eliminar otros tags HTML
        texto = re.sub(r'<(?!/?[bi]>)[^>]+>', '', texto)
        return texto.strip()

    # =========================================================================
    # MÉTODOS PRIVADOS - DOCX
    # =========================================================================

    def _configurar_estilos_docx(self, doc):
        """Configura los estilos del documento Word."""
        from docx.shared import Pt, RGBColor
        from docx.enum.style import WD_STYLE_TYPE

        styles = doc.styles

        # Estilo Heading 1
        h1 = styles['Heading 1']
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(26, 54, 93)

        # Estilo Heading 2
        h2 = styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(44, 82, 130)

        # Estilo Normal
        normal = styles['Normal']
        normal.font.size = Pt(11)
        normal.font.name = 'Arial'

    def _agregar_portada_docx(self, doc, proyecto: Proyecto, documento: DocumentoEIA):
        """Agrega la portada al documento Word."""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Espacio superior
        for _ in range(5):
            doc.add_paragraph()

        # Título principal
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("ESTUDIO DE IMPACTO AMBIENTAL")
        run.bold = True
        run.font.size = Pt(24)

        # Subtítulo
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitulo.add_run("Conforme a la Ley 19.300 y D.S. 40/2012")
        run.font.size = Pt(14)

        doc.add_paragraph()
        doc.add_paragraph()

        # Nombre del proyecto
        nombre = doc.add_paragraph()
        nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = nombre.add_run(proyecto.nombre or "Proyecto")
        run.bold = True
        run.font.size = Pt(20)

        doc.add_paragraph()
        doc.add_paragraph()

        # Información del proyecto
        metadatos = documento.metadatos or {}

        info_lines = [
            f"Titular: {proyecto.titular or 'N/A'}",
            f"Ubicación: {proyecto.region or 'N/A'}, {proyecto.comuna or 'N/A'}",
        ]

        if metadatos.get('empresa_consultora'):
            info_lines.append(f"Consultora: {metadatos['empresa_consultora']}")

        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)

        doc.add_paragraph()
        doc.add_paragraph()

        # Fecha y versión
        fecha = doc.add_paragraph()
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_str = metadatos.get('fecha_elaboracion', datetime.now().strftime('%d/%m/%Y'))
        fecha.add_run(f"Fecha: {fecha_str}")

        version = doc.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version.add_run(f"Versión: v{documento.version}")

    def _agregar_indice_docx(self, doc, documento: DocumentoEIA):
        """Agrega el índice al documento Word."""
        from docx.shared import Pt

        # Título del índice
        titulo = doc.add_heading("ÍNDICE DE CONTENIDOS", level=1)

        # Agregar entradas del índice
        capitulos = documento.contenido_capitulos or {}

        for num_str, capitulo_data in sorted(capitulos.items(), key=lambda x: int(x[0])):
            num = int(num_str)
            titulo_cap = capitulo_data.get('titulo', f'Capítulo {num}')

            p = doc.add_paragraph()
            run = p.add_run(f"{num}. {titulo_cap}")
            run.bold = True

            # Agregar subsecciones si existen
            subsecciones = capitulo_data.get('subsecciones', {})
            seccion_num = 1
            for seccion_titulo in subsecciones.keys():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                p.add_run(f"{num}.{seccion_num} {seccion_titulo}")
                seccion_num += 1

    def _agregar_capitulo_docx(self, doc, numero: int, capitulo_data: Dict[str, Any]):
        """Agrega un capítulo al documento Word."""
        from docx.shared import Pt

        titulo = capitulo_data.get('titulo', f'Capítulo {numero}')
        contenido = capitulo_data.get('contenido', '')

        # Agregar título del capítulo
        doc.add_heading(f"Capítulo {numero}: {titulo}", level=1)

        # Parsear markdown básico y agregar contenido
        self._agregar_contenido_markdown_docx(doc, contenido)

        # Salto de página después de cada capítulo
        doc.add_page_break()

    def _agregar_contenido_markdown_docx(self, doc, contenido_md: str):
        """Convierte contenido Markdown básico a formato Word."""
        from docx.shared import Pt

        lineas = contenido_md.split('\n')
        i = 0

        while i < len(lineas):
            linea = lineas[i].rstrip()

            # Heading 2: ## Título
            if linea.startswith('## '):
                doc.add_heading(linea[3:], level=2)

            # Heading 3: ### Título
            elif linea.startswith('### '):
                doc.add_heading(linea[4:], level=3)

            # Heading 4: #### Título
            elif linea.startswith('#### '):
                doc.add_heading(linea[5:], level=4)

            # Lista con viñetas: - item o * item
            elif linea.startswith('- ') or linea.startswith('* '):
                p = doc.add_paragraph(linea[2:], style='List Bullet')

            # Lista numerada: 1. item
            elif re.match(r'^\d+\. ', linea):
                texto = re.sub(r'^\d+\. ', '', linea)
                p = doc.add_paragraph(texto, style='List Number')

            # Cita: > texto
            elif linea.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(36)
                run = p.add_run(linea[2:])
                run.italic = True

            # Línea vacía
            elif not linea.strip():
                pass  # Saltar líneas vacías

            # Párrafo normal
            else:
                # Procesar negritas y cursivas básicas
                texto = self._procesar_formato_inline(linea)
                if texto.strip():
                    p = doc.add_paragraph(texto)

            i += 1

    def _procesar_formato_inline(self, texto: str) -> str:
        """Procesa formato inline básico (negritas, cursivas)."""
        # Por ahora solo limpiamos las marcas, en una implementación
        # más completa se aplicarían los estilos correspondientes
        texto = re.sub(r'\*\*(.+?)\*\*', r'\1', texto)  # Negritas
        texto = re.sub(r'\*(.+?)\*', r'\1', texto)      # Cursivas
        texto = re.sub(r'`(.+?)`', r'\1', texto)        # Código inline
        return texto

    # =========================================================================
    # MÉTODOS PRIVADOS - e-SEIA XML
    # =========================================================================

    def _construir_xml_eseia(
        self,
        documento: DocumentoEIA,
        proyecto: Proyecto,
        config: ConfiguracionESEIA
    ) -> Any:
        """Construye la estructura XML para e-SEIA."""
        from lxml import etree

        # Namespace e-SEIA (simplificado - ajustar según esquema real)
        NSMAP = {
            None: "http://www.sea.gob.cl/eseia/v2",
            'xsi': "http://www.w3.org/2001/XMLSchema-instance"
        }

        # Root element
        root = etree.Element("EIA", nsmap=NSMAP)
        root.set("{http://www.w3.org/2001/XMLSchema-instance}schemaLocation",
                 "http://www.sea.gob.cl/eseia/v2 eseia_v2.xsd")

        # Metadatos del documento
        metadatos_elem = etree.SubElement(root, "Metadatos")
        etree.SubElement(metadatos_elem, "Version").text = str(documento.version)
        etree.SubElement(metadatos_elem, "FechaGeneracion").text = datetime.now().isoformat()
        etree.SubElement(metadatos_elem, "VersionEsquema").text = config.version_esquema

        # Información del proyecto
        proyecto_elem = etree.SubElement(root, "Proyecto")
        etree.SubElement(proyecto_elem, "Nombre").text = proyecto.nombre
        etree.SubElement(proyecto_elem, "Titular").text = proyecto.titular or ""
        etree.SubElement(proyecto_elem, "Region").text = proyecto.region or ""
        etree.SubElement(proyecto_elem, "Comuna").text = proyecto.comuna or ""

        if proyecto.superficie_ha:
            etree.SubElement(proyecto_elem, "SuperficieHa").text = str(proyecto.superficie_ha)
        if proyecto.inversion_musd:
            etree.SubElement(proyecto_elem, "InversionMUSD").text = str(proyecto.inversion_musd)

        # Contenido por capítulos
        contenido_elem = etree.SubElement(root, "Contenido")
        capitulos = documento.contenido_capitulos or {}

        for num_str, capitulo_data in sorted(capitulos.items(), key=lambda x: int(x[0])):
            capitulo_elem = etree.SubElement(contenido_elem, "Capitulo")
            capitulo_elem.set("numero", num_str)

            etree.SubElement(capitulo_elem, "Titulo").text = capitulo_data.get('titulo', '')

            # Contenido como CDATA
            texto_elem = etree.SubElement(capitulo_elem, "Texto")
            texto_elem.text = etree.CDATA(capitulo_data.get('contenido', ''))

            # Subsecciones
            subsecciones = capitulo_data.get('subsecciones', {})
            if subsecciones:
                secciones_elem = etree.SubElement(capitulo_elem, "Secciones")
                for seccion_titulo, seccion_contenido in subsecciones.items():
                    seccion_elem = etree.SubElement(secciones_elem, "Seccion")
                    etree.SubElement(seccion_elem, "Titulo").text = seccion_titulo
                    contenido_sec = etree.SubElement(seccion_elem, "Contenido")
                    contenido_sec.text = etree.CDATA(seccion_contenido if isinstance(seccion_contenido, str) else "")

        # Anexos (placeholder para futura implementación)
        if config.incluir_anexos_digitales:
            anexos_elem = etree.SubElement(root, "Anexos")
            # TODO: Implementar anexos digitales

        return root

    # =========================================================================
    # MÉTODOS PRIVADOS - UTILIDADES
    # =========================================================================

    async def _get_documento(self, db: AsyncSession, documento_id: int) -> DocumentoEIA:
        """Obtiene un documento EIA por ID."""
        query = select(DocumentoEIA).where(DocumentoEIA.id == documento_id)
        result = await db.execute(query)
        documento = result.scalar_one_or_none()

        if not documento:
            raise ValueError(f"Documento EIA {documento_id} no encontrado")

        return documento

    async def _get_proyecto(self, db: AsyncSession, proyecto_id: int) -> Proyecto:
        """Obtiene un proyecto por ID."""
        query = select(Proyecto).where(Proyecto.id == proyecto_id)
        result = await db.execute(query)
        proyecto = result.scalar_one_or_none()

        if not proyecto:
            raise ValueError(f"Proyecto {proyecto_id} no encontrado")

        return proyecto

    def _generar_nombre_archivo(
        self,
        proyecto: Proyecto,
        documento: DocumentoEIA,
        extension: str
    ) -> str:
        """Genera un nombre de archivo único para la exportación."""
        # Limpiar nombre del proyecto
        nombre_limpio = re.sub(r'[^\w\s-]', '', proyecto.nombre or 'proyecto')
        nombre_limpio = re.sub(r'\s+', '_', nombre_limpio)[:50]

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        return f"EIA_{nombre_limpio}_v{documento.version}_{timestamp}.{extension}"

    async def _registrar_exportacion(
        self,
        db: AsyncSession,
        documento_id: int,
        formato: FormatoExportacion,
        filepath: Optional[str],
        tamano: int,
        config: Dict[str, Any],
        error: Optional[str] = None
    ) -> ExportacionEIA:
        """Registra una exportación en la base de datos."""
        exportacion = ExportacionEIA(
            documento_id=documento_id,
            formato=formato.value,
            archivo_path=filepath,
            tamano_bytes=tamano if tamano > 0 else None,
            generado_exitoso=error is None,
            error_mensaje=error,
            configuracion=config
        )

        db.add(exportacion)
        await db.commit()
        await db.refresh(exportacion)

        return exportacion
